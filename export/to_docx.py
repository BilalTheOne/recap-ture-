from docx import Document

from export.to_markdown import _format_timestamp


def export_docx(lines: list[dict], path: str) -> None:
    document = Document()
    for line in lines:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(
            f"[{_format_timestamp(line['start'])}] {line['speaker']}:"
        )
        run.bold = True
        document.add_paragraph(line["text"])
    document.save(path)
